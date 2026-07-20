"""
Motor de calificación asistida por IA, usando la API de Gemini (Google).

Nota para quien mantenga este código: este módulo requiere el paquete
`google-genai` (NO `google-generativeai`, que está descontinuado), `pydantic`
y `python-docx` (para leer archivos .docx). Instalar con:
pip install google-genai pydantic python-docx

Requiere una variable de entorno GEMINI_API_KEY (o pasar api_key explícito),
obtenida gratis en https://aistudio.google.com/apikey

Este es el ÚNICO módulo del proyecto que depende de pydantic/google-genai.
Recibe y devuelve los tipos definidos en `schemas.py` (dataclasses puras),
para que el resto de la app se pueda probar sin estas dependencias.
"""

from __future__ import annotations

import io
import os
import time

from pydantic import BaseModel, Field
from google import genai
from google.genai import types

from .prompt_builder import construir_instrucciones
from .schemas import ArchivoReferencia, DesgloseItem, EjemploCalificado, ResultadoCalificacion

# gemini-3.5-flash: modelo actual, multimodal (imagen + PDF nativo), gratuito
# en el free tier a la fecha de escritura (verificar en
# https://ai.google.dev/gemini-api/docs/pricing si cambia).
MODEL = "gemini-3.5-flash"

MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def extraer_texto_docx(data: bytes) -> str:
    """Gemini no interpreta archivos .docx de forma nativa (a diferencia de
    PDF e imágenes), así que si un .docx llega como material de referencia,
    ejemplo, o incluso como el propio informe a calificar (Carla mencionó
    que a veces los informes son Word), extraemos su texto localmente y se
    lo pasamos como texto plano en vez de mandarle el archivo crudo."""
    import docx  # python-docx

    doc = docx.Document(io.BytesIO(data))
    partes = [p.text for p in doc.paragraphs if p.text.strip()]
    for tabla in doc.tables:
        for fila in tabla.rows:
            partes.append(" | ".join(c.text.strip() for c in fila.cells))
    return "\n".join(partes)


def _como_contenido_gemini(data: bytes, mime_type: str, etiqueta: str) -> list:
    """Devuelve la(s) parte(s) para mandarle a Gemini. Imagen o PDF: se manda
    tal cual (Gemini los entiende de forma nativa). Word (.docx): se extrae
    el texto localmente y se manda como texto plano."""
    if mime_type == MIME_DOCX:
        texto = extraer_texto_docx(data)
        return [f"--- Contenido de '{etiqueta}' (Word, texto extraído) ---\n{texto}"]
    return [types.Part.from_bytes(data=data, mime_type=mime_type)]


# ---- Esquema interno solo para pedirle a Gemini salida estructurada ----
# (se convierte a schemas.ResultadoCalificacion antes de devolver nada)

class _DesgloseItemSchema(BaseModel):
    criterio: str
    puntaje_obtenido: float
    puntaje_maximo: float
    comentario: str
    pagina: int = Field(
        default=0,
        description="Número de página del PDF (empezando en 1) donde está esa sección. "
        "Usa 0 si no aplica o no estás segura -- no adivines.",
    )


class _ResultadoSchema(BaseModel):
    nombre_detectado: str = Field(
        description="Nombre del estudiante (o integrantes del grupo separados por "
        "coma, si es informe) tal como se identifica en el documento"
    )
    confianza_nombre: str = Field(description="'alta', 'media' o 'baja'")
    grupo_detectado: str = Field(default="", description="Número de grupo si aplica, o vacío")
    practica_detectada: str = Field(
        description="De qué práctica/tema trata realmente el contenido del documento, según lo "
        "que ves (aunque sea distinto al nombre que te dieron). Sé breve, ej: 'Ley de Stefan-Boltzmann'."
    )
    nota_sugerida: float
    desglose: list[_DesgloseItemSchema]
    observaciones: str


def _cliente(api_key: str | None = None) -> genai.Client:
    key = api_key or os.environ.get("GEMINI_API_KEY")
    if not key:
        raise RuntimeError(
            "Falta la API key de Gemini. Configúrala como variable de entorno "
            "GEMINI_API_KEY (en Streamlit Cloud: Settings -> Secrets)."
        )
    return genai.Client(api_key=key)


def _es_error_de_limite(error: Exception) -> bool:
    """El SDK no siempre expone un tipo de excepción específico y estable
    para 429 -- se revisa el mensaje/código, que sí es consistente."""
    texto = str(error).lower()
    return any(
        marca in texto
        for marca in ("429", "resource_exhausted", "rate limit", "quota", "too many requests")
    )


def _llamar_con_reintentos(client: genai.Client, contents: list, intentos: int = 4):
    """Reintenta con espera creciente (5s, 15s, 30s...) si Gemini responde
    que se alcanzó el límite de solicitudes por minuto -- para que un lote
    de varios documentos no se corte a la mitad por esto."""
    espera = 5
    ultimo_error = None
    for intento in range(intentos):
        try:
            return client.models.generate_content(
                model=MODEL,
                contents=contents,
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                    response_schema=_ResultadoSchema,
                    temperature=0.2,  # priorizamos consistencia sobre creatividad
                ),
            )
        except Exception as e:
            ultimo_error = e
            if not _es_error_de_limite(e) or intento == intentos - 1:
                raise
            time.sleep(espera)
            espera = min(espera * 2, 60)
    raise ultimo_error  # pragma: no cover -- inalcanzable, por claridad


def calificar(
    tipo: str,  # "coloquio" | "informe"
    nombre_practica: str,
    rubrica_texto: str,
    puntaje_maximo: float,
    contenido: bytes,
    contenido_mime: str,  # "image/jpeg" | "image/png" | "application/pdf" | docx
    nombres_candidatos: list[str] | None = None,
    ejemplos: list[EjemploCalificado] | None = None,
    material_archivos: list[ArchivoReferencia] | None = None,
    api_key: str | None = None,
) -> ResultadoCalificacion:
    """Envía el material de referencia + ejemplos + el documento nuevo a
    Gemini y devuelve una sugerencia de calificación estructurada. NO escribe
    nada en el Excel; eso lo hace el módulo de escritura después de que
    Carla revise."""

    nombres_candidatos = nombres_candidatos or []
    ejemplos = ejemplos or []
    material_archivos = material_archivos or []

    # Los .docx de referencia no se mandan como archivo -- se convierten a
    # texto y se fusionan con el texto libre que haya escrito Carla.
    texto_extra_docx = []
    archivos_nativos = []
    for arch in material_archivos:
        if arch.mime_type == MIME_DOCX:
            texto_extra_docx.append(
                f"--- Contenido de '{arch.nombre or 'documento Word'}' ---\n"
                + extraer_texto_docx(arch.data)
            )
        else:
            archivos_nativos.append(arch)

    rubrica_texto_final = rubrica_texto
    if texto_extra_docx:
        rubrica_texto_final = (rubrica_texto + "\n\n" + "\n\n".join(texto_extra_docx)).strip()

    instrucciones = construir_instrucciones(
        tipo=tipo,
        nombre_practica=nombre_practica,
        rubrica_texto=rubrica_texto_final,
        puntaje_maximo=puntaje_maximo,
        nombres_candidatos=nombres_candidatos,
        num_ejemplos=len(ejemplos),
        num_archivos_referencia=len(archivos_nativos),
    )

    contents: list = [instrucciones]

    for arch in archivos_nativos:
        contents.append(f"--- Material de referencia: {arch.nombre or arch.mime_type} ---")
        contents.append(types.Part.from_bytes(data=arch.data, mime_type=arch.mime_type))

    for i, ej in enumerate(ejemplos, start=1):
        contents.append(
            f"--- Ejemplo {i}: ya calificado por Carla con nota real "
            f"{ej.nota_asignada}/{puntaje_maximo} ---"
        )
        contents.extend(_como_contenido_gemini(ej.data, ej.mime_type, f"ejemplo {i}"))
        if ej.comentario:
            contents.append(f"Comentario de Carla sobre este ejemplo: {ej.comentario}")

    contents.append("--- Documento NUEVO a calificar ahora ---")
    contents.extend(_como_contenido_gemini(contenido, contenido_mime, "documento nuevo"))

    client = _cliente(api_key)
    response = _llamar_con_reintentos(client, contents)
    r = _ResultadoSchema.model_validate_json(response.text)

    return ResultadoCalificacion(
        nombre_detectado=r.nombre_detectado,
        confianza_nombre=r.confianza_nombre,
        nota_sugerida=r.nota_sugerida,
        desglose=[
            DesgloseItem(
                criterio=d.criterio,
                puntaje_obtenido=d.puntaje_obtenido,
                puntaje_maximo=d.puntaje_maximo,
                comentario=d.comentario,
                pagina=d.pagina if d.pagina > 0 else None,
            )
            for d in r.desglose
        ],
        observaciones=r.observaciones,
        grupo_detectado=r.grupo_detectado or None,
        practica_detectada=r.practica_detectada,
    )
